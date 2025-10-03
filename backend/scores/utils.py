from __future__ import annotations

from django.db.models import Sum, Count, Q, Avg, Max, Min
from django.core.exceptions import ValidationError
from django.db import transaction
from django.conf import settings
from django.utils import timezone
from decimal import Decimal
import io
import os
import zipfile
import statistics
import tempfile
from datetime import datetime

import pandas as pd

from .models import Score, TestResult, CommentTemplate
from schools.models import School
from students.models import Student
from tests.models import TestDefinition, QuestionGroup

def calculate_test_results(student, test, force_temporary=False):
    """学生のテスト結果を計算"""
    from django.utils import timezone
    
    # スコアの合計を計算
    scores = Score.objects.filter(student=student, test=test)
    total_score = scores.aggregate(total=Sum('score'))['total'] or 0
    
    # 正答率を計算
    max_possible_score = test.max_score
    correct_rate = (total_score / max_possible_score * 100) if max_possible_score > 0 else 0
    
    # 締切状況を確認
    is_deadline_passed = timezone.now() > test.schedule.deadline_at
    is_final_calculation = is_deadline_passed and not force_temporary
    
    # 順位計算
    school_rank, school_total = calculate_school_rank_enhanced(student, test, total_score, is_final_calculation)
    national_rank, national_total = calculate_national_rank_enhanced(test, total_score, is_final_calculation)
    
    # コメント生成
    comment = generate_comment(student.classroom.school, test.subject, total_score)
    
    # TestResultを作成または更新
    defaults = {
        'total_score': total_score,
        'correct_rate': correct_rate,
        'comment': comment,
    }
    
    if is_final_calculation:
        # 確定後の順位更新
        defaults.update({
            'school_rank_final': school_rank,
            'national_rank_final': national_rank,
            'school_total_final': school_total,
            'national_total_final': national_total,
            'is_rank_finalized': True,
            'rank_finalized_at': timezone.now(),
        })
    else:
        # 一時的な順位更新
        defaults.update({
            'school_rank_temporary': school_rank,
            'national_rank_temporary': national_rank,
            'school_total_temporary': school_total,
            'national_total_temporary': national_total,
        })
    
    result, created = TestResult.objects.update_or_create(
        student=student,
        test=test,
        defaults=defaults
    )
    
    return result

def calculate_rankings_unified(student, test, total_score, is_final=False):
    """統一された順位計算ロジック"""
    from django.db.models import Count

    school = student.classroom.school if student.classroom else None

    # 出席者のみを対象とした基本クエリ
    base_queryset = TestResult.objects.filter(
        test=test,
        student__scores__test=test,
        student__scores__attendance=True
    ).distinct()

    # 塾内順位
    school_rank, school_total = None, 0
    if school:
        school_queryset = base_queryset.filter(student__classroom__school=school)
        school_rank = school_queryset.filter(total_score__gt=total_score).count() + 1
        school_total = school_queryset.count()

        # 自分が新規の場合は総数に追加
        if not school_queryset.filter(student=student).exists():
            school_total += 1

    # 学年順位
    grade_queryset = base_queryset.filter(student__grade=student.grade)
    grade_rank = grade_queryset.filter(total_score__gt=total_score).count() + 1
    grade_total = grade_queryset.count()

    # 自分が新規の場合は総数に追加
    if not grade_queryset.filter(student=student).exists():
        grade_total += 1

    # 全国順位
    national_rank = base_queryset.filter(total_score__gt=total_score).count() + 1
    national_total = base_queryset.count()

    # 自分が新規の場合は総数に追加
    if not base_queryset.filter(student=student).exists():
        national_total += 1

    return {
        'school_rank': school_rank,
        'school_total': school_total,
        'grade_rank': grade_rank,
        'grade_total': grade_total,
        'national_rank': national_rank,
        'national_total': national_total
    }

def calculate_school_rank_enhanced(student, test, total_score, is_final=False):
    """塾内順位を計算（拡張版：一時的・確定後に対応）"""
    rankings = calculate_rankings_unified(student, test, total_score, is_final)
    return rankings['school_rank'], rankings['school_total']

def calculate_school_rank(student, test, total_score):
    """塾内順位を計算（後方互換性のため維持）"""
    return calculate_school_rank_enhanced(student, test, total_score, is_final=False)

def calculate_national_rank_enhanced(test, total_score, is_final=False):
    """全国順位を計算（拡張版：一時的・確定後に対応）"""
    # ダミー学生オブジェクトを使って統一ロジックを呼び出し
    from students.models import Student
    dummy_student = Student(grade='1', classroom=None)
    rankings = calculate_rankings_unified(dummy_student, test, total_score, is_final)
    return rankings['national_rank'], rankings['national_total']

def calculate_national_rank(test, total_score):
    """全国順位を計算（後方互換性のため維持）"""
    return calculate_national_rank_enhanced(test, total_score, is_final=False)

def calculate_grade_rank_enhanced(student, test, total_score, is_final=False):
    """学年順位を計算（拡張版：一時的・確定後に対応）"""
    rankings = calculate_rankings_unified(student, test, total_score, is_final)
    return rankings['grade_rank'], rankings['grade_total']

def calculate_grade_rank(student, test, total_score):
    """学年順位を計算"""
    return calculate_grade_rank_enhanced(student, test, total_score, is_final=False)

def convert_student_grade_to_test_grade_level(student_grade):
    """学生の学年文字列をTestDefinitionのgrade_levelに変換"""
    if not student_grade:
        return None
    
    try:
        # 数字のみの場合
        grade_num = int(student_grade)
        if 1 <= grade_num <= 6:
            return f'elementary_{grade_num}'
        elif grade_num == 7:
            return 'middle_1'
        elif grade_num == 8:
            return 'middle_2'
        elif grade_num == 9:
            return 'middle_3'
    except (ValueError, TypeError):
        # 数字変換に失敗した場合、文字列から判定
        if '小' in student_grade:
            try:
                grade_num = int(student_grade.replace('小', '').replace('年', '').replace('生', ''))
                if 1 <= grade_num <= 6:
                    return f'elementary_{grade_num}'
            except:
                pass
        elif '中' in student_grade:
            try:
                grade_num = int(student_grade.replace('中', '').replace('年', '').replace('生', ''))
                if 1 <= grade_num <= 3:
                    return f'middle_{grade_num}'
            except:
                pass
    
    return None

def calculate_grade_statistics(student, test):
    """学年ごとの統計情報（平均点・偏差値）を計算"""
    import numpy as np
    from django.db.models import Avg
    
    # 同じ学年のテスト結果を取得
    grade_results = TestResult.objects.filter(
        test=test,
        student__grade=student.grade
    )
    
    if not grade_results.exists():
        return {
            'average': 0,
            'deviation_score': 50,  # デフォルト偏差値
            'participant_count': 0
        }
    
    # 平均点計算
    average = grade_results.aggregate(Avg('total_score'))['total_score__avg'] or 0
    
    # 偏差値計算用のスコアリスト取得
    scores = list(grade_results.values_list('total_score', flat=True))
    
    if len(scores) <= 1:
        deviation_score = 50  # 1人以下の場合はデフォルト偏差値
    else:
        scores_array = np.array(scores)
        mean_score = np.mean(scores_array)
        std_score = np.std(scores_array, ddof=1)  # 標本標準偏差
        
        if std_score == 0:
            deviation_score = 50  # 標準偏差が0の場合はデフォルト偏差値
        else:
            # 学生の成績を取得
            student_result = grade_results.filter(student=student).first()
            if student_result:
                student_score = student_result.total_score
                deviation_score = 50 + (student_score - mean_score) / std_score * 10
            else:
                deviation_score = 50
    
    return {
        'average': round(average, 1),
        'deviation_score': round(deviation_score, 1),
        'participant_count': len(scores)
    }

def generate_comment(school, subject, score):
    """点数に応じたコメントを生成"""
    # まず塾専用のコメントテンプレートを探す
    template = CommentTemplate.objects.filter(
        school=school,
        subject=subject,
        score_range_min__lte=score,
        score_range_max__gte=score,
        is_active=True
    ).first()
    
    # 塾専用がない場合はデフォルトテンプレートを使用
    if not template:
        template = CommentTemplate.objects.filter(
            school=None,
            subject=subject,
            score_range_min__lte=score,
            score_range_max__gte=score,
            is_active=True,
            is_default=True
        ).first()
    
    return template.template_text if template else "よく頑張りました。"

# SchoolStatistics機能は削除されました

def bulk_calculate_test_results(test, force_recalculate=False):
    """指定されたテストの全学生の結果を一括計算・更新（最適化版）"""
    from django.utils import timezone
    from django.db import transaction
    from django.db.models import Sum, F
    from students.models import Student

    print(f"=== 一括集計開始: {test} ===")

    with transaction.atomic():
        # 出席者のスコアを一括取得
        student_scores = Score.objects.filter(
            test=test,
            attendance=True
        ).values('student').annotate(
            total_score=Sum('score')
        ).select_related('student')

        # 学生別の合計点辞書を作成
        score_dict = {}
        students_data = []

        for score_data in student_scores:
            student_id = score_data['student']
            total_score = score_data['total_score'] or 0

            try:
                student = Student.objects.get(id=student_id)
                score_dict[student_id] = total_score

                # 正答率計算
                correct_rate = (total_score / test.max_score * 100) if test.max_score > 0 else 0

                students_data.append({
                    'student': student,
                    'total_score': total_score,
                    'correct_rate': correct_rate
                })

            except Student.DoesNotExist:
                continue

        print(f"対象学生数: {len(students_data)}")

        # 一括順位計算
        rankings_data = calculate_bulk_rankings(students_data, test)

        # TestResult一括更新
        test_results_to_update = []
        test_results_to_create = []

        for student_data in students_data:
            student = student_data['student']
            total_score = student_data['total_score']
            correct_rate = student_data['correct_rate']

            # 順位データを取得
            rankings = rankings_data.get(student.id, {})

            # コメント生成
            comment = generate_comment(
                student.classroom.school if student.classroom else None,
                test.subject,
                total_score
            )

            # 締切状況確認
            is_deadline_passed = timezone.now() > test.schedule.deadline_at

            # TestResultの既存レコードを確認
            try:
                test_result = TestResult.objects.get(student=student, test=test)
                # 既存レコードを更新
                test_result.total_score = total_score
                test_result.correct_rate = correct_rate
                test_result.comment = comment

                if is_deadline_passed:
                    # 確定後の順位
                    test_result.school_rank_final = rankings.get('school_rank')
                    test_result.national_rank_final = rankings.get('national_rank')
                    test_result.grade_rank = rankings.get('grade_rank')
                    test_result.school_total_final = rankings.get('school_total')
                    test_result.national_total_final = rankings.get('national_total')
                    test_result.grade_total = rankings.get('grade_total')
                    test_result.is_rank_finalized = True
                    test_result.rank_finalized_at = timezone.now()
                else:
                    # 一時的順位
                    test_result.school_rank_temporary = rankings.get('school_rank')
                    test_result.national_rank_temporary = rankings.get('national_rank')
                    test_result.grade_rank = rankings.get('grade_rank')
                    test_result.school_total_temporary = rankings.get('school_total')
                    test_result.national_total_temporary = rankings.get('national_total')
                    test_result.grade_total = rankings.get('grade_total')

                test_results_to_update.append(test_result)

            except TestResult.DoesNotExist:
                # 新規レコードを作成
                test_result_data = {
                    'student': student,
                    'test': test,
                    'total_score': total_score,
                    'correct_rate': correct_rate,
                    'comment': comment,
                    'grade_rank': rankings.get('grade_rank'),
                    'grade_total': rankings.get('grade_total')
                }

                if is_deadline_passed:
                    test_result_data.update({
                        'school_rank_final': rankings.get('school_rank'),
                        'national_rank_final': rankings.get('national_rank'),
                        'school_total_final': rankings.get('school_total'),
                        'national_total_final': rankings.get('national_total'),
                        'is_rank_finalized': True,
                        'rank_finalized_at': timezone.now()
                    })
                else:
                    test_result_data.update({
                        'school_rank_temporary': rankings.get('school_rank'),
                        'national_rank_temporary': rankings.get('national_rank'),
                        'school_total_temporary': rankings.get('school_total'),
                        'national_total_temporary': rankings.get('national_total')
                    })

                test_results_to_create.append(TestResult(**test_result_data))

        # 一括更新・作成実行
        if test_results_to_update:
            TestResult.objects.bulk_update(
                test_results_to_update,
                ['total_score', 'correct_rate', 'comment', 'grade_rank', 'grade_total',
                 'school_rank_temporary', 'national_rank_temporary',
                 'school_total_temporary', 'national_total_temporary',
                 'school_rank_final', 'national_rank_final',
                 'school_total_final', 'national_total_final',
                 'is_rank_finalized', 'rank_finalized_at']
            )
            print(f"更新済み: {len(test_results_to_update)}件")

        if test_results_to_create:
            TestResult.objects.bulk_create(test_results_to_create)
            print(f"新規作成: {len(test_results_to_create)}件")

    print(f"=== 一括集計完了: {test} ===")
    return len(students_data)

def calculate_bulk_rankings(students_data, test):
    """学生データから一括で順位を計算"""
    rankings_data = {}

    # 学生を点数順にソート
    sorted_students = sorted(students_data, key=lambda x: x['total_score'], reverse=True)

    # 塾別グループ化
    school_groups = {}
    grade_groups = {}

    for student_data in students_data:
        student = student_data['student']
        school = student.classroom.school if student.classroom else None
        grade = student.grade

        # 塾別グループ
        if school:
            if school.id not in school_groups:
                school_groups[school.id] = []
            school_groups[school.id].append(student_data)

        # 学年別グループ
        if grade not in grade_groups:
            grade_groups[grade] = []
        grade_groups[grade].append(student_data)

    # 全国順位計算
    for i, student_data in enumerate(sorted_students):
        student = student_data['student']
        total_score = student_data['total_score']

        # 全国順位（同点は同順位）
        national_rank = 1
        for other_data in sorted_students:
            if other_data['total_score'] > total_score:
                national_rank += 1
            else:
                break

        rankings_data[student.id] = {
            'national_rank': national_rank,
            'national_total': len(sorted_students)
        }

    # 塾別順位計算
    for school_id, school_students in school_groups.items():
        school_sorted = sorted(school_students, key=lambda x: x['total_score'], reverse=True)

        for student_data in school_students:
            student = student_data['student']
            total_score = student_data['total_score']

            school_rank = 1
            for other_data in school_sorted:
                if other_data['total_score'] > total_score:
                    school_rank += 1
                else:
                    break

            rankings_data[student.id]['school_rank'] = school_rank
            rankings_data[student.id]['school_total'] = len(school_sorted)

    # 学年順位計算
    for grade, grade_students in grade_groups.items():
        grade_sorted = sorted(grade_students, key=lambda x: x['total_score'], reverse=True)

        for student_data in grade_students:
            student = student_data['student']
            total_score = student_data['total_score']

            grade_rank = 1
            for other_data in grade_sorted:
                if other_data['total_score'] > total_score:
                    grade_rank += 1
                else:
                    break

            rankings_data[student.id]['grade_rank'] = grade_rank
            rankings_data[student.id]['grade_total'] = len(grade_sorted)

    return rankings_data


def set_cell_background(cell, rgb_color):
    """python-docxのセル背景色を設定"""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    try:
        fill = str(rgb_color.rgb)
    except AttributeError:
        fill = 'FFFFFF'
    shading.set(qn('w:fill'), fill)
    tc_pr.append(shading)


def create_beautiful_word_report(report_data):
    """
    美しいデザインのWordレポートを生成
    A3横向きサイズで見やすいレイアウト
    """
    try:
        import tempfile
        import os
        from datetime import datetime
        try:
            from docx import Document
            from docx.shared import Inches, Cm, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.shared import OxmlElement, qn
            from docx.enum.section import WD_SECTION
            from docx.enum.section import WD_ORIENT
        except ImportError as import_error:
            print(f"python-docxのインポートエラー: {import_error}")
            # フォールバック: 簡易的なテキストファイルを生成
            return create_fallback_text_report(report_data)
        
        # 新しいドキュメントを作成
        doc = Document()
        
        # A3横向きに設定
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(42)  # A3横向き幅
        section.page_height = Cm(29.7)  # A3横向き高さ
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        
        # ヘッダー部分
        header = doc.add_heading('', level=0)
        header_run = header.runs[0] if header.runs else header.add_run()
        header_run.text = '個別成績表'
        header_run.font.size = Pt(24)
        header_run.font.color.rgb = RGBColor(31, 73, 125)  # 濃い青色
        header_run.bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生徒情報テーブル
        student_info = report_data['student_info']
        test_info = report_data['test_info']
        
        info_table = doc.add_table(rows=2, cols=4)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # ヘッダー行
        info_cells = info_table.rows[0].cells
        info_cells[0].text = '生徒氏名'
        info_cells[1].text = '生徒ID'
        info_cells[2].text = '学年・教室'
        info_cells[3].text = 'テスト実施日'
        
        # データ行
        data_cells = info_table.rows[1].cells
        data_cells[0].text = student_info['name']
        data_cells[1].text = student_info['id']
        data_cells[2].text = f"{student_info['grade']} {student_info['classroom_name']}"
        data_cells[3].text = str(test_info.get('date', ''))
        
        # テーブルのスタイル設定
        for row in info_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
        
        # ヘッダー行の背景色設定
        for cell in info_table.rows[0].cells:
            set_cell_background(cell, RGBColor(189, 215, 238))  # 薄い青色
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()  # 空行
        
        # 科目別結果
        subjects = report_data.get('subjects', {})
        
        if subjects:
            # 科目を分類（左：算数/数学、右：国語/英語）
            math_subjects = []
            language_subjects = []
            
            for subject_code, subject_data in subjects.items():
                if 'math' in subject_code.lower() or '算数' in subject_data.get('name', '') or '数学' in subject_data.get('name', ''):
                    math_subjects.append((subject_code, subject_data))
                else:  # 国語、英語など
                    language_subjects.append((subject_code, subject_data))
            
            # メインテーブル（2列）
            main_table = doc.add_table(rows=1, cols=2)
            main_table.style = 'Table Grid'
            
            # テーブル幅を調整
            for column in main_table.columns:
                for cell in column.cells:
                    cell.width = Cm(18)  # 各列の幅を18cmに設定
            
            # 左側セル：算数/数学
            left_cell = main_table.cell(0, 0)
            left_paragraph = left_cell.paragraphs[0]
            left_run = left_paragraph.add_run('算数・数学')
            left_run.font.size = Pt(16)
            left_run.font.color.rgb = RGBColor(31, 73, 125)
            left_run.bold = True
            left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if math_subjects:
                for subject_code, subject_data in math_subjects:
                    create_subject_section(left_cell, subject_data)
            else:
                # 学年に応じたサンプルデータを表示
                student_info = report_data.get('student_info', {})
                grade = student_info.get('grade', '不明')
                is_elementary = grade.startswith('小') or grade.isdigit()
                
                sample_math_data = {
                    'name': f"{'算数' if is_elementary else '数学'}（{grade}）",
                    'total_score': '-',
                    'max_score': 100,
                    'average_score': '-',
                    'national_rank': '-',
                    'total_students': '-',
                    'comment': 'テスト結果データがありません。'
                }
                create_subject_section(left_cell, sample_math_data)
            
            # 右側セル：国語/英語
            right_cell = main_table.cell(0, 1)
            right_paragraph = right_cell.paragraphs[0]
            right_run = right_paragraph.add_run('国語・英語')
            right_run.font.size = Pt(16)
            right_run.font.color.rgb = RGBColor(31, 73, 125)
            right_run.bold = True
            right_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if language_subjects:
                for subject_code, subject_data in language_subjects:
                    create_subject_section(right_cell, subject_data)
            else:
                # 学年に応じたサンプルデータを表示
                student_info = report_data.get('student_info', {})
                grade = student_info.get('grade', '不明')
                is_elementary = grade.startswith('小') or grade.isdigit()
                
                sample_language_data = {
                    'name': f"{'国語' if is_elementary else '英語'}（{grade}）",
                    'total_score': '-',
                    'max_score': 100,
                    'average_score': '-',
                    'national_rank': '-',
                    'total_students': '-',
                    'comment': 'テスト結果データがありません。'
                }
                create_subject_section(right_cell, sample_language_data)
            
            # セルの背景色を設定（薄いグレー）
            set_cell_background(left_cell, RGBColor(248, 248, 248))
            set_cell_background(right_cell, RGBColor(252, 252, 252))
        
        # 総合コメント欄
        doc.add_paragraph()
        comment_header = doc.add_heading('総合所見', level=2)
        comment_header.runs[0].font.color.rgb = RGBColor(31, 73, 125)
        
        # コメントボックス
        comment_table = doc.add_table(rows=1, cols=1)
        comment_table.style = 'Table Grid'
        comment_cell = comment_table.cell(0, 0)
        
        # セルに余白を設定
        comment_cell._tc.get_or_add_tcPr().append(
            OxmlElement('w:tcMar')
        )
        
        # 各科目のコメントを統合
        all_comments = []
        for subject_data in subjects.values():
            if subject_data.get('comment'):
                all_comments.append(f"【{subject_data['name']}】{subject_data['comment']}")
        
        comment_text = '\n\n'.join(all_comments) if all_comments else "よく頑張りました。引き続き学習を継続してください。"
        
        # コメントを段落として追加
        comment_paragraph = comment_cell.paragraphs[0]
        comment_run = comment_paragraph.add_run(comment_text)
        comment_run.font.size = Pt(11)
        comment_run.font.name = 'ＭＳ Ｐゴシック'  # 日本語フォント
        comment_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # コメント欄の背景色を設定
        set_cell_background(comment_cell, RGBColor(252, 252, 252))  # 非常に薄いグレー
        
        # フッター
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(f"作成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # メディアディレクトリに保存
        from django.conf import settings
        
        media_reports_dir = os.path.join(settings.MEDIA_ROOT, 'reports')
        os.makedirs(media_reports_dir, exist_ok=True)
        
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(media_reports_dir, filename)
        doc.save(file_path)
        
        # ファイル権限を設定（読み取り可能にする）
        import stat
        os.chmod(file_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IROTH)  # 644
        
        print(f"Word帳票生成完了: {file_path}")
        
        return file_path
        
    except Exception as e:
        print(f"Word帳票生成エラー: {str(e)}")
        import traceback
        traceback.print_exc()
        raise e

def create_bar_chart(doc, subject_data):
    """
    得点と平均点の比較棒グラフを生成してWordに挿入
    """
    try:
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')
        import numpy as np
        import os
        from django.conf import settings
        from docx.shared import Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 日本語フォント設定
        matplotlib.rcParams['font.family'] = ['DejaVu Sans', 'Hiragino Sans', 'Yu Gothic', 'Meiryo', 'Takao', 'IPAexGothic', 'IPAPGothic', 'VL PGothic', 'Noto Sans CJK JP']
        
        # データ準備
        categories = ['全国', '塾内']
        my_scores = [subject_data['total_score'], subject_data['total_score']]
        averages = [subject_data['average_score'], subject_data['school_average']]
        
        x = np.arange(len(categories))
        width = 0.35
        
        # グラフ作成
        fig, ax = plt.subplots(figsize=(8, 5))
        bars1 = ax.bar(x - width/2, my_scores, width, label='あなたの得点', color='#1f73b5')
        bars2 = ax.bar(x + width/2, averages, width, label='平均点', color='#ff7f0e')
        
        # 軸とラベル設定
        ax.set_xlabel('比較対象')
        ax.set_ylabel('得点')
        ax.set_title(f'{subject_data["name"]} 得点比較')
        ax.set_xticks(x)
        ax.set_xticklabels(categories)
        ax.legend()
        ax.grid(True, alpha=0.3)
        
        # 棒の上に数値表示
        for bar in bars1:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                   f'{height:.0f}', ha='center', va='bottom')
        
        for bar in bars2:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                   f'{height:.1f}', ha='center', va='bottom')
        
        # 保存
        media_reports_dir = os.path.join(settings.MEDIA_ROOT, 'reports')
        os.makedirs(media_reports_dir, exist_ok=True)
        chart_path = os.path.join(media_reports_dir, f'bar_{subject_data["name"]}_{np.random.randint(1000, 9999)}.png')
        
        plt.tight_layout()
        plt.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        # Wordに挿入
        paragraph = doc.add_paragraph()
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.add_picture(chart_path, width=Cm(12))  # チャートサイズ縮小
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 一時ファイル削除
        try:
            os.remove(chart_path)
        except:
            pass
            
    except Exception as e:
        print(f"棒グラフ生成エラー: {str(e)}")

def generate_word_report_old(report_data):
    """Word形式の成績表を生成（A3横向き二つ折り）"""
    try:
        from docx import Document
        from docx.shared import Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import tempfile
        import os
        
        # 新しいドキュメント作成
        doc = Document()
        
        # ページ設定（A3横向き）
        section = doc.sections[0]
        section.page_width = Inches(16.54)  # A3横
        section.page_height = Inches(11.69)  # A3縦
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # タイトル
        title = doc.add_heading('全国学力診断テスト 成績表', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生徒情報ヘッダー
        info_table = doc.add_table(rows=2, cols=4)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 生徒情報入力
        cells = info_table.rows[0].cells
        cells[0].text = f"氏名: {report_data['student_info']['name']}"
        cells[1].text = f"学年: {report_data['student_info']['grade']}"
        cells[2].text = f"学校: {report_data['student_info']['school_name']}"
        cells[3].text = f"クラス: {report_data['student_info']['classroom_name']}"
        
        cells = info_table.rows[1].cells
        cells[0].text = f"生徒ID: {report_data['student_info']['id']}"
        cells[1].text = f"実施年度: {report_data['test_info']['year']}年度"
        cells[2].text = f"実施期間: {report_data['test_info']['period']}"
        cells[3].text = f"実施日: {report_data['test_info']['date']}"
        
        doc.add_paragraph()  # 空行
        
        # 2列レイアウト用のテーブル（左右分割）
        main_table = doc.add_table(rows=1, cols=2)
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 左側と右側のセル
        left_cell = main_table.rows[0].cells[0]
        right_cell = main_table.rows[0].cells[1]
        
        # 教科データを左右に分割
        subjects_list = list(report_data['subjects'].items())
        
        # 左側（1つ目の教科）
        if len(subjects_list) > 0:
            subject_key, subject_data = subjects_list[0]
            add_subject_content(left_cell, subject_data)
        
        # 右側（2つ目の教科）
        if len(subjects_list) > 1:
            subject_key, subject_data = subjects_list[1]
            add_subject_content(right_cell, subject_data)
        
        # ファイル保存
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        # メディアディレクトリに移動
        from django.conf import settings
        from datetime import datetime
        reports_dir = os.path.join(settings.MEDIA_ROOT, 'reports')
        os.makedirs(reports_dir, exist_ok=True)
        
        final_path = os.path.join(reports_dir, f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        os.rename(temp_file.name, final_path)
        
        return final_path
        
    except Exception as e:
        print(f"Word帳票生成エラー: {str(e)}")
        raise e

def get_grade_level_from_student_grade(student_grade):
    """生徒の学年から対応するテスト学年レベルを取得"""
    # 数値・文字列の両方に対応
    grade_mapping = {
        # 数値形式
        1: 'elementary_1', '1': 'elementary_1',
        2: 'elementary_2', '2': 'elementary_2',
        3: 'elementary_3', '3': 'elementary_3',
        4: 'elementary_4', '4': 'elementary_4',
        5: 'elementary_5', '5': 'elementary_5',
        6: 'elementary_6', '6': 'elementary_6',
        7: 'middle_1', '7': 'middle_1',
        8: 'middle_2', '8': 'middle_2',
        9: 'middle_3', '9': 'middle_3',
        # 文字列形式
        '小1': 'elementary_1',
        '小2': 'elementary_2',
        '小3': 'elementary_3',
        '小4': 'elementary_4',
        '小5': 'elementary_5',
        '小6': 'elementary_6',
        '中1': 'middle_1',
        '中2': 'middle_2',
        '中3': 'middle_3'
    }
    return grade_mapping.get(student_grade)

def is_student_test_grade_match(student, test):
    """生徒の学年とテストの対象学年が一致するかチェック"""
    expected_grade_level = get_grade_level_from_student_grade(student.grade)
    return expected_grade_level == test.grade_level

def import_scores_from_excel(file_path, year, period, subject=None, grade_level=None, school_id=None):
    """
    Excelファイルから得点データを一括インポート（新形式対応）
    塾ID・塾名・教室ID・教室名・生徒ID・生徒名・学年・年度・期間・出席・大問・合計点
    """
    try:
        # CSVファイルを読み込み（BOM対応）
        try:
            df = pd.read_csv(file_path, encoding='utf-8-sig')  # BOM付きUTF-8
        except:
            try:
                df = pd.read_csv(file_path, encoding='utf-8')  # 通常のUTF-8
            except:
                try:
                    df = pd.read_csv(file_path, encoding='shift_jis')  # Shift_JIS
                except:
                    df = pd.read_excel(file_path)  # Excel形式
        
        # 列名を正規化（空白文字除去）
        df.columns = df.columns.str.strip()
        
        # 統合テンプレート形式か単一教科形式かを判定
        is_unified_template = subject is None or any(col for col in df.columns if '_大問' in col)
        
        if not is_unified_template:
            # 単一教科の場合のテスト構造を取得
            structure = get_test_template_structure(year, period, subject, grade_level)
            if not structure:
                period_display = {'spring': '春季', 'summer': '夏季', 'winter': '冬季'}.get(period, period)
                subject_display = {'japanese': '国語', 'math': '算数', 'english': '英語', 'mathematics': '数学'}.get(subject, subject)
                grade_display = {'elementary': '小学生', 'middle_school': '中学生'}.get(grade_level, grade_level) if grade_level else ''
                raise ValidationError(f"{year}年度{period_display}{grade_display}{subject_display}テストが見つかりません")
            
            test = structure['test']
        else:
            # 統合テンプレートの場合はtestをNoneに設定（後で各教科別に処理）
            test = None
            structure = None
        
        # 基本列をチェック
        expected_columns = ['塾ID', '塾名', '教室ID', '教室名', '生徒ID', '生徒名', '学年', '年度', '期間', '出席']
        missing_columns = [col for col in expected_columns if col not in df.columns]
        if missing_columns:
            raise ValidationError(f"必要な列が不足しています: {', '.join(missing_columns)}\\n実際の列: {list(df.columns)}")
        
        # 統合テンプレート形式の検出（デバッグ用）
        unified_columns = [col for col in df.columns if '_大問' in col]
        single_columns = [col for col in df.columns if col.startswith('大問')]
        print(f"列の検出: 統合形式列={unified_columns}, 単一形式列={single_columns}")
        
        # 空行を除去
        df = df.dropna(how='all')
        
        # データ型を文字列に統一し、数値フィールドの小数点を除去
        for col in df.columns:
            df[col] = df[col].astype(str).str.strip()
            # 数値IDフィールドの小数点を除去
            if col in ['塾ID', '教室ID', '生徒ID']:
                df[col] = df[col].str.replace('.0', '', regex=False)
                # ExcelエラーやNaN値をクリーンアップ
                df[col] = df[col].replace(['nan', '#N/A', '#REF!', '#VALUE!', '#DIV/0!', '#NAME?', '#NULL!', '#NUM!'], '')
        
        created_scores = []
        updated_scores = []
        errors = []
        
        with transaction.atomic():
            for index, row in df.iterrows():
                try:
                    # 空行をスキップ
                    if pd.isna(row['生徒ID']) or str(row['生徒ID']).strip() == '' or str(row['生徒ID']).strip() == 'nan':
                        continue
                    
                    student_id = str(row['生徒ID']).strip()
                    
                    # 必須フィールドの検証
                    if str(row['塾ID']).strip() == '' or str(row['塾ID']).strip() == 'nan':
                        errors.append(f'行 {index + 2}: 塾IDが空です')
                        continue
                    
                    if str(row['教室ID']).strip() == '' or str(row['教室ID']).strip() == 'nan':
                        errors.append(f'行 {index + 2}: 教室IDが空です (塾ID: {row["塾ID"]})')
                        continue
                    
                    # 教室を取得
                    from classrooms.models import Classroom
                    classroom = Classroom.objects.filter(
                        classroom_id=row['教室ID'],
                        school__school_id=row['塾ID']
                    ).first()
                    
                    if not classroom:
                        errors.append(f'行 {index + 2}: 教室が見つかりません (塾ID: {row["塾ID"]}, 教室ID: {row["教室ID"]})')
                        continue
                    
                    # 生徒の存在確認・作成（StudentEnrollment経由）
                    from students.models import StudentEnrollment
                    
                    # まず該当期間の受講生徒を探す
                    enrollment = StudentEnrollment.objects.filter(
                        student__student_id=student_id,
                        student__classroom=classroom,
                        year=year,
                        period=period,
                        is_active=True
                    ).first()
                    
                    if enrollment:
                        student = enrollment.student
                    else:
                        # 生徒が存在するかチェック（他の期間含む）
                        existing_student = Student.objects.filter(
                            student_id=student_id,
                            classroom=classroom
                        ).first()
                        
                        if existing_student:
                            # 既存生徒の新期間登録を作成
                            student = existing_student
                            enrollment = StudentEnrollment.objects.create(
                                student=student,
                                year=year,
                                period=period,
                                is_active=True
                            )
                            print(f"新期間受講登録作成: 生徒ID {student_id} ({student.name}) - {year}年度{period}期")
                        else:
                            # 完全に新しい生徒を作成
                            student_name = row.get('生徒名', f'生徒{student_id}')
                            student_grade = row.get('学年', '1')
                            
                            # 学年形式を数値に変換
                            if isinstance(student_grade, str):
                                if student_grade.startswith('小'):
                                    student_grade = student_grade.replace('小', '')
                                elif student_grade.startswith('中'):
                                    middle_grade = int(student_grade.replace('中', ''))
                                    student_grade = str(middle_grade + 6)
                            
                            student = Student.objects.create(
                                student_id=student_id,
                                name=student_name,
                                grade=student_grade,
                                classroom=classroom,
                                is_active=True
                            )
                            
                            # 新生徒の受講登録も作成
                            enrollment = StudentEnrollment.objects.create(
                                student=student,
                                year=year,
                                period=period,
                                is_active=True
                            )
                            print(f"新生徒・受講登録作成: 生徒ID {student_id} ({student_name}) - {year}年度{period}期")
                    
                    # 出席状況を処理
                    attendance_value = str(row.get('出席', '出席')).strip()
                    is_attendance = attendance_value in ['出席', '○', '1', 'True', 'true']
                    
                    # 各大問の得点を処理（単一教科・統合テンプレート両方対応）
                    
                    # 統合テンプレートかどうかを判定（教科名_大問の形式の列があるかチェック）
                    is_unified_template = any(col for col in df.columns if '_大問' in col)
                    print(f"行{index + 2}: 統合テンプレート判定: {is_unified_template}, 生徒: {student.name}, 学年: {student.grade}")
                    
                    if is_unified_template:
                        # 統合テンプレートの場合：複数教科を処理
                        
                        # 利用可能な教科とテスト構造を取得
                        from tests.models import TestDefinition
                        available_tests = TestDefinition.objects.filter(
                            schedule__year=year,
                            schedule__period=period,
                            is_active=True
                        )
                        print(f"利用可能なテスト: {[(str(t), t.subject, t.grade_level) for t in available_tests]}")
                        
                        # 教科マッピング
                        subject_mapping = {
                            '国語': 'japanese',
                            '算数': 'math',
                            '英語': 'english',
                            '数学': 'mathematics'
                        }
                        
                        for subject_display, subject_code in subject_mapping.items():
                            # この生徒の学年に対応する教科かチェック
                            is_elementary = student.grade and student.grade.startswith('小')
                            
                            # 小学生は国語・算数、中学生は英語・数学
                            if is_elementary and subject_code in ['japanese', 'math']:
                                should_process = True
                            elif not is_elementary and subject_code in ['english', 'mathematics']:
                                should_process = True
                            else:
                                should_process = False
                            
                            if not should_process:
                                continue
                            
                            # 対応するテストを探す
                            test_for_subject = available_tests.filter(
                                subject=subject_code,
                                grade_level__startswith='elementary_' if is_elementary else 'middle_'
                            ).first()
                            
                            if not test_for_subject:
                                print(f"警告: {subject_display}のテストが見つかりません (生徒: {student.name}, 学年: {student.grade})")
                                continue
                            
                            # この教科の大問を処理
                            subject_structure = get_test_template_structure(year, period, subject_code, test_for_subject.grade_level)
                            if not subject_structure:
                                print(f"警告: {subject_display}のテスト構造が見つかりません")
                                continue
                            
                            print(f"処理中: {subject_display} (テスト: {str(test_for_subject)})")
                            
                            for group_info in subject_structure['question_groups']:
                                column_name = f"{subject_display}_大問{group_info['group_number']}"
                                
                                # 得点値を取得
                                score_value = None
                                if column_name in df.columns:
                                    score_value = row[column_name]
                                
                                if pd.isna(score_value) or str(score_value).strip() in ['', 'nan']:
                                    continue  # 空の場合はスキップ
                                
                                try:
                                    score_value = float(score_value)
                                    if score_value < 0 or score_value > group_info['max_score']:
                                        errors.append(f"行{index + 2}: {column_name}の得点が範囲外です (0-{group_info['max_score']})")
                                        continue
                                except (ValueError, TypeError):
                                    errors.append(f"行{index + 2}: {column_name}の得点が無効です")
                                    continue
                                
                                # 生徒の学年とテストの対象学年をチェック
                                if not is_student_test_grade_match(student, test_for_subject):
                                    expected_grade_level = get_grade_level_from_student_grade(student.grade)
                                    errors.append(f"行{index + 2}: {student.name}の学年({student.grade})とテスト対象学年({test_for_subject.grade_level})が一致しません。期待される学年レベル: {expected_grade_level}")
                                    continue
                                
                                # 得点レコードを作成または更新
                                try:
                                    group = QuestionGroup.objects.get(
                                        test=test_for_subject,
                                        group_number=group_info['group_number']
                                    )
                                    score_obj, created = Score.objects.update_or_create(
                                        student=student,
                                        test=test_for_subject,
                                        question_group=group,
                                        defaults={
                                            'score': score_value,
                                            'attendance': is_attendance
                                        }
                                    )
                                    
                                    if created:
                                        created_scores.append(score_obj)
                                    else:
                                        updated_scores.append(score_obj)
                                    
                                    print(f"成功: {student.name} - {subject_display} 大問{group_info['group_number']}: {score_value}点")
                                except QuestionGroup.DoesNotExist:
                                    errors.append(f"行{index + 2}: {column_name}の大問が見つかりません (テスト: {str(test_for_subject)})")
                                    continue
                    else:
                        # 単一教科テンプレートの場合：従来の処理
                        for group_info in structure['question_groups']:
                            group = QuestionGroup.objects.get(id=group_info['id'])
                            column_name = f"大問{group_info['group_number']}"
                            
                            # 得点値を取得
                            score_value = None
                            if column_name in df.columns:
                                score_value = row[column_name]
                            
                            if pd.isna(score_value) or str(score_value).strip() in ['', 'nan']:
                                continue  # 空の場合はスキップ
                            
                            try:
                                score_value = float(score_value)
                                if score_value < 0 or score_value > group_info['max_score']:
                                    errors.append(f"行{index + 2}: {column_name}の得点が範囲外です (0-{group_info['max_score']})")
                                    continue
                            except (ValueError, TypeError):
                                errors.append(f"行{index + 2}: {column_name}の得点が無効です")
                                continue
                            
                            # 生徒の学年とテストの対象学年をチェック
                            if not is_student_test_grade_match(student, test):
                                expected_grade_level = get_grade_level_from_student_grade(student.grade)
                                errors.append(f"行{index + 2}: {student.name}の学年({student.grade})とテスト対象学年({test.grade_level})が一致しません。期待される学年レベル: {expected_grade_level}")
                                continue
                            
                            # 得点レコードを作成または更新
                            score_obj, created = Score.objects.update_or_create(
                                student=student,
                                test=test,
                                question_group=group,
                                defaults={
                                    'score': score_value,
                                    'attendance': is_attendance
                                }
                            )
                            
                            if created:
                                created_scores.append(score_obj)
                            else:
                                updated_scores.append(score_obj)
                    
                except Exception as e:
                    student_info = f"生徒ID: {student_id}" if 'student_id' in locals() else "不明な生徒"
                    errors.append(f"行{index + 2} ({student_info}): {str(e)}")
                    print(f"インポートエラー - 行{index + 2}: {student_info} - {str(e)}")
        
        period_display = {'spring': '春季', 'summer': '夏季', 'winter': '冬季'}.get(period, period)
        
        # インポート詳細情報を追加
        import_details = []
        for score in created_scores + updated_scores:
            import_details.append({
                'student_name': score.student.name,
                'student_id': score.student.student_id,
                'subject': score.test.get_subject_display(),
                'question_group': f"大問{score.question_group.group_number}",
                'score': score.score,
                'attendance': score.attendance
            })
        
        return {
            'success': True,
            'created_scores': len(created_scores),
            'updated_scores': len(updated_scores),
            'test_info': f"{year}年度{period_display} 統合テンプレート",
            'import_details': import_details[:20],  # 最初の20件のみ表示
            'total_processed': len(import_details),
            'errors': errors
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def generate_unified_score_template(year, period, grade_level):
    """
    指定された学年のすべての教科を含む統合テンプレートを生成
    小学生: 国語・算数, 中学生: 英語・数学
    """
    # 学年レベルに応じて教科を決定
    if grade_level == 'elementary':
        subjects = ['japanese', 'math']  # 国語、算数
    elif grade_level == 'middle_school':
        subjects = ['english', 'mathematics']  # 英語、数学
    else:
        raise ValidationError(f"無効な学年レベルです: {grade_level}")
    
    # 各教科のテスト構造を取得
    all_structures = {}
    for subject in subjects:
        structure = get_test_template_structure(year, period, subject, grade_level)
        if structure:
            all_structures[subject] = structure
    
    if not all_structures:
        period_display = {'spring': '春季', 'summer': '夏季', 'winter': '冬季'}.get(period, period)
        grade_display = {'elementary': '小学生', 'middle_school': '中学生'}.get(grade_level, grade_level)
        raise ValidationError(f"{year}年度{period_display}{grade_display}のテストが見つかりません")
    
    # 基本情報列（生徒登録と同じ形式）
    columns = [
        '塾ID', '塾名', '教室ID', '教室名', 
        '生徒ID', '生徒名', '学年', '年度', '期間',
        '出席'  # 出席状況
    ]
    
    # 各教科の大問ごとの列を追加
    subject_columns = {}
    for subject, structure in all_structures.items():
        subject_display = {'japanese': '国語', 'math': '算数', 'english': '英語', 'mathematics': '数学'}.get(subject, subject)
        subject_columns[subject] = []
        
        for group in structure['question_groups']:
            column_name = f"{subject_display}_大問{group['group_number']}"
            columns.append(column_name)
            subject_columns[subject].append(column_name)
        
        # 教科別合計点列
        total_column = f"{subject_display}_合計点"
        columns.append(total_column)
        subject_columns[subject].append(total_column)
    
    # 全体合計点列
    columns.append('全体合計点')
    
    # サンプルデータ
    sample_data = {col: [] for col in columns}
    
    # 期間表示を変換
    period_display_jp = {'spring': '春期', 'summer': '夏期', 'winter': '冬期'}.get(period, period)
    
    # サンプル行を3つ作成
    sample_names = ['田中太郎', '佐藤花子', '鈴木次郎']
    if grade_level == 'elementary':
        sample_grades = ['小6', '小5', '小4']
    else:
        sample_grades = ['中1', '中2', '中3']
    
    for i in range(3):
        sample_data['塾ID'].append('100001')
        sample_data['塾名'].append('サンプル学習塾')
        sample_data['教室ID'].append('001001')
        sample_data['教室名'].append('メイン教室')
        sample_data['生徒ID'].append(f'{123456 + i}')
        sample_data['生徒名'].append(sample_names[i])
        sample_data['学年'].append(sample_grades[i])
        sample_data['年度'].append(str(year))
        sample_data['期間'].append(period_display_jp)
        sample_data['出席'].append('出席')
        
        overall_total = 0
        
        # 各教科のサンプル点数を設定
        for subject, structure in all_structures.items():
            subject_total = 0
            
            for group in structure['question_groups']:
                subject_display = {'japanese': '国語', 'math': '算数', 'english': '英語', 'mathematics': '数学'}.get(subject, subject)
                column_name = f"{subject_display}_大問{group['group_number']}"
                
                # サンプル点数（満点から徐々に減らす）
                score = max(0, group['max_score'] - i * 2)
                sample_data[column_name].append(score)
                subject_total += score
            
            # 教科別合計点
            subject_display = {'japanese': '国語', 'math': '算数', 'english': '英語', 'mathematics': '数学'}.get(subject, subject)
            total_column = f"{subject_display}_合計点"
            sample_data[total_column].append(subject_total)
            overall_total += subject_total
        
        # 全体合計点
        sample_data['全体合計点'].append(overall_total)
    
    # 空行を追加（データ入力用）
    for _ in range(3):
        for col in columns:
            if col in ['年度', '期間']:
                sample_data[col].append(str(year) if col == '年度' else period_display_jp)
            else:
                sample_data[col].append('')
    
    df = pd.DataFrame(sample_data)
    return df, all_structures

def get_available_tests():
    """
    利用可能なテスト一覧を取得
    """
    tests = TestDefinition.objects.select_related('schedule').order_by('-schedule__year', '-schedule__period', 'subject')
    
    available_tests = []
    for test in tests:
        period_display = {'spring': '春季', 'summer': '夏季', 'winter': '冬季'}.get(test.schedule.period, test.schedule.period)
        available_tests.append({
            'year': test.schedule.year,
            'period': test.schedule.period,
            'period_display': period_display,
            'subject': test.subject,
            'subject_display': test.get_subject_display(),
            'name': str(test),
            'question_groups_count': test.question_groups.count()
        })
    
    return available_tests

def get_test_summary_by_school_type(year, period, school_type):
    """
    学校種別での集計結果を取得
    """
    try:
        # 簡易実装 - 集計データの取得
        return {
            'success': True,
            'test_summary': {
                'year': year,
                'period': period,
                'school_type': school_type,
                'total_students': 0,
                'average_score': 0.0,
                'average_correct_rate': 0.0,
            },
            'school_summaries': []
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def recalculate_test_results_for_test(test, grade=None):
    """
    指定されたテストの結果を再計算（学年指定可能）
    """
    from students.models import Student
    
    # 対象学生を取得
    students_query = Student.objects.filter(
        scores__test=test
    ).distinct()
    
    if grade:
        students_query = students_query.filter(grade=grade)
    
    results = []
    for student in students_query:
        result = calculate_test_results(student, test)
        results.append(result)
    
    return results

def get_test_summary(year, period, subject, grade_level=None):
    """
    保存されたテスト集計結果を取得
    """
    from .models import TestSummary, SchoolTestSummary
    from tests.models import TestDefinition
    
    try:
        query = TestDefinition.objects.filter(
            schedule__year=year,
            schedule__period=period,
            subject=subject
        )
        
        if grade_level:
            query = query.filter(grade_level=grade_level)
        
        test = query.first()
        if not test:
            return {'success': False, 'error': 'テストが見つかりません'}
        
        test_summary = TestSummary.objects.get(test=test)
        school_summaries = SchoolTestSummary.objects.filter(
            test_summary=test_summary
        ).select_related('school').order_by('rank_among_schools')
        
        return {
            'test_summary': test_summary,
            'school_summaries': school_summaries,
            'success': True
        }
        
    except (TestDefinition.DoesNotExist, TestSummary.DoesNotExist):
        return {'success': False, 'error': '集計結果が見つかりません'}
    except Exception as e:
        return {'success': False, 'error': str(e)}


# =============================================================
# 個別成績表PDF生成ユーティリティ
# =============================================================

PERIOD_LABELS = {
    'spring': '春期',
    'summer': '夏期',
    'winter': '冬期',
}

PERIOD_ITERATIONS = {
    'spring': '第1回',
    'summer': '第2回',
    'winter': '第3回',
}

PERIOD_SHORT_LABELS = {
    'spring': '春',
    'summer': '夏',
    'winter': '冬',
}

PERIOD_ORDER = {'spring': 1, 'summer': 2, 'winter': 3}

SUBJECT_DISPLAY = {
    'math': '算数',
    'mathematics': '数学',
    'japanese': '国語',
    'english': '英語',
}

SUBJECT_ORDER = {
    'math': 1,
    'mathematics': 1,
    'japanese': 2,
    'english': 3,
}

REPORTS_SUBDIR = 'reports'
PDF_FONTS_REGISTERED = False


def _ensure_reports_dir() -> str:
    """帳票保存先ディレクトリを作成してパスを返す"""
    reports_dir = os.path.join(settings.MEDIA_ROOT, REPORTS_SUBDIR)
    os.makedirs(reports_dir, exist_ok=True)
    return reports_dir


def _build_download_url(abs_path: str) -> str:
    """媒体ファイルの絶対パスからダウンロードURLを生成"""
    rel_path = os.path.relpath(abs_path, settings.MEDIA_ROOT)
    if rel_path.startswith('..'):
        return abs_path  # MEDIA_ROOT外に保存されている場合は絶対パスを返す
    rel_path = rel_path.replace(os.sep, '/')
    base = settings.MEDIA_URL
    if not base.endswith('/'):
        base = f"{base}/"
    return f"{base}{rel_path}"


def _period_display(period: str) -> str:
    return PERIOD_LABELS.get(period, period)


def _iteration_display(period: str) -> str:
    return PERIOD_ITERATIONS.get(period, '')


def _short_period_label(period: str) -> str:
    return PERIOD_SHORT_LABELS.get(period, period)


def _subject_display(subject: str) -> str:
    if subject in SUBJECT_DISPLAY:
        return SUBJECT_DISPLAY[subject]
    subject_dict = dict(TestDefinition.SUBJECTS)
    return subject_dict.get(subject, subject)


def _calculate_deviation(score: float, population_scores: list[float]) -> float | None:
    if not population_scores:
        return None
    if len(population_scores) == 1:
        return 50.0
    mean_value = statistics.mean(population_scores)
    std_dev = statistics.pstdev(population_scores)
    if std_dev == 0:
        return 50.0
    deviation = 50 + 10 * ((score - mean_value) / std_dev)
    return round(deviation, 1)


def _register_pdf_fonts() -> None:
    """ReportLabで日本語フォントを利用できるよう登録"""
    global PDF_FONTS_REGISTERED
    if PDF_FONTS_REGISTERED:
        return
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    except ImportError:
        return
    pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))
    pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMincho-W3'))
    PDF_FONTS_REGISTERED = True


def _calculate_combined_average(grade: str, year: int, period: str) -> float:
    grade_totals = TestResult.objects.filter(
        test__schedule__year=year,
        test__schedule__period=period,
        student__grade=grade
    ).values('student').annotate(total=Sum('total_score'))
    if not grade_totals:
        return 0.0
    total_sum = sum(item['total'] for item in grade_totals)
    return float(total_sum / len(grade_totals))


def _calculate_combined_metrics(student: Student, year: int, period: str, total_score: float) -> dict:
    base_qs = TestResult.objects.filter(
        test__schedule__year=year,
        test__schedule__period=period
    )

    grade_qs = list(
        base_qs.filter(student__grade=student.grade)
        .values('student', 'student__classroom__school_id')
        .annotate(total=Sum('total_score'))
    )

    grade_scores = [item['total'] for item in grade_qs]
    grade_rank = (sum(1 for item in grade_qs if item['total'] > total_score) + 1) if grade_qs else None
    grade_total = len(grade_qs)
    grade_average = (sum(grade_scores) / grade_total) if grade_total else 0
    grade_deviation = _calculate_deviation(total_score, grade_scores)

    school_rank = None
    school_total = 0
    school_average = 0
    school_scores = []
    school_id = student.classroom.school_id if student.classroom and student.classroom.school else None
    if school_id:
        school_entries = [item for item in grade_qs if item['student__classroom__school_id'] == school_id]
        school_scores = [item['total'] for item in school_entries]
        school_total = len(school_entries)
        if school_total:
            school_rank = sum(1 for item in school_entries if item['total'] > total_score) + 1
            school_average = sum(school_scores) / school_total

    national_qs = list(
        base_qs.values('student').annotate(total=Sum('total_score'))
    )
    national_scores = [item['total'] for item in national_qs]
    national_rank = (sum(1 for item in national_qs if item['total'] > total_score) + 1) if national_qs else None
    national_total = len(national_qs)
    national_average = (sum(national_scores) / national_total) if national_total else 0
    national_deviation = _calculate_deviation(total_score, national_scores)

    return {
        'grade_rank': grade_rank,
        'grade_total': grade_total,
        'grade_average': round(grade_average, 1) if grade_total else 0,
        'grade_deviation': grade_deviation,
        'school_rank': school_rank,
        'school_total': school_total,
        'school_average': round(school_average, 1) if school_total else 0,
        'school_deviation': _calculate_deviation(total_score, school_scores) if school_scores else None,
        'national_rank': national_rank,
        'national_total': national_total,
        'national_average': round(national_average, 1) if national_total else 0,
        'national_deviation': national_deviation,
    }


def _collect_trend_data(student: Student, grade_level: str) -> dict:
    if not grade_level:
        return {'overall': [], 'subjects': {}}

    results_qs = TestResult.objects.filter(
        student=student,
        test__grade_level=grade_level
    ).select_related('test__schedule').order_by('test__schedule__year', 'test__schedule__period')

    timeline: dict[tuple[int, str], dict] = {}
    for result in results_qs:
        schedule = result.test.schedule
        key = (schedule.year, schedule.period)
        entry = timeline.setdefault(key, {
            'label': f"{str(schedule.year)[2:]}{_short_period_label(schedule.period)}",
            'subjects': {},
        })
        entry['subjects'][result.test.subject] = result.total_score

    sorted_keys = sorted(timeline.keys(), key=lambda x: (x[0], PERIOD_ORDER.get(x[1], 9)))

    overall_trend = []
    subject_trend: dict[str, list] = {}

    for year_value, period_value in sorted_keys:
        entry = timeline[(year_value, period_value)]
        subject_scores = [score for score in entry['subjects'].values() if score is not None]
        total_score = sum(subject_scores)
        overall_average = _calculate_combined_average(student.grade, year_value, period_value)
        overall_trend.append({
            'label': entry['label'],
            'score': total_score,
            'average': overall_average,
        })

        for subject_code, score in entry['subjects'].items():
            subject_avg = TestResult.objects.filter(
                test__schedule__year=year_value,
                test__schedule__period=period_value,
                test__subject=subject_code,
                student__grade=student.grade
            ).aggregate(avg=Avg('total_score'))['avg'] or 0
            subject_trend.setdefault(subject_code, [])
            subject_trend[subject_code].append({
                'label': entry['label'],
                'score': score,
                'average': float(subject_avg),
            })

    return {'overall': overall_trend, 'subjects': subject_trend}


def _collect_individual_report_data(student_id: str, year: int, period: str) -> tuple[dict | None, str | None]:
    student = Student.objects.select_related('classroom__school').filter(student_id=student_id).first()
    if not student:
        return None, '対象の生徒が見つかりません'

    grade_level = get_grade_level_from_student_grade(student.grade)

    tests_qs = TestDefinition.objects.filter(
        schedule__year=year,
        schedule__period=period
    )
    if grade_level:
        tests_qs = tests_qs.filter(grade_level=grade_level)

    tests = list(tests_qs.select_related('schedule').prefetch_related('question_groups'))
    if not tests:
        return None, '指定条件のテストが見つかりません'

    results_map = {
        result.test_id: result
        for result in TestResult.objects.filter(student=student, test__in=tests)
        .select_related('test', 'test__schedule')
    }
    if not results_map:
        return None, '指定条件の成績が登録されていません'

    schedule = tests[0].schedule
    test_date = schedule.actual_date or schedule.planned_date

    student_info = {
        'id': student.student_id,
        'name': student.name,
        'grade': student.grade,
        'school_name': student.classroom.school.name if student.classroom and student.classroom.school else '',
        'classroom_name': student.classroom.name if student.classroom else '',
        'membership_type': student.classroom.school.get_membership_type_display() if student.classroom and student.classroom.school else '',
    }

    test_info = {
        'year': int(year),
        'period': period,
        'period_display': _period_display(period),
        'iteration': _iteration_display(period),
        'date': test_date.strftime('%Y.%m.%d') if test_date else '',
        'grade_level': grade_level,
    }

    subject_entries = []
    subjects_data = {}

    for test in tests:
        result = results_map.get(test.id)
        subject_code = test.subject
        subject_name = _subject_display(subject_code)

        if not result:
            continue

        student_scores_qs = Score.objects.filter(student=student, test=test).select_related('question_group')
        student_scores = {score.question_group_id: score for score in student_scores_qs}
        attended = any(score.attendance for score in student_scores_qs)

        group_averages = {
            item['question_group_id']: item['avg']
            for item in Score.objects.filter(test=test)
            .values('question_group_id')
            .annotate(avg=Avg('score'))
        }

        subject_results_qs = TestResult.objects.filter(test=test)
        national_average = subject_results_qs.aggregate(avg=Avg('total_score'))['avg'] or 0
        national_total = subject_results_qs.count()

        school_average = 0
        school_high = None
        school_low = None
        if student.classroom and student.classroom.school:
            school_results = subject_results_qs.filter(student__classroom__school=student.classroom.school)
            if school_results.exists():
                school_average = school_results.aggregate(avg=Avg('total_score'))['avg'] or 0
                school_high = school_results.aggregate(max=Max('total_score'))['max']
                school_low = school_results.aggregate(min=Min('total_score'))['min']

        national_rank, national_total_rank = result.get_current_national_rank()
        school_rank, school_total_rank = result.get_current_school_rank()

        grade_rank = result.grade_rank
        grade_total = result.grade_total

        deviation = float(result.grade_deviation_score) if result.grade_deviation_score is not None else None
        if deviation is None:
            population_scores = list(
                subject_results_qs.filter(student__grade=student.grade).values_list('total_score', flat=True)
            )
            deviation = _calculate_deviation(result.total_score, population_scores)

        question_details = []
        for group in test.question_groups.order_by('group_number'):
            score_obj = student_scores.get(group.id)
            score_value = score_obj.score if score_obj else None
            max_score = group.max_score
            national_avg = group_averages.get(group.id, 0)
            question_details.append({
                'number': group.group_number,
                'title': group.title,
                'score': score_value,
                'max_score': max_score,
                'national_average': national_avg,
                'correct_rate': round((score_value / max_score) * 100, 1) if score_value is not None and max_score else None,
                'national_correct_rate': round((national_avg / max_score) * 100, 1) if max_score else None,
            })

        subjects_data[subject_code] = {
            'code': subject_code,
            'name': subject_name,
            'total_score': result.total_score,
            'max_score': test.max_score,
            'deviation': deviation,
            'attendance': attended,
            'rankings': {
                'national': {'rank': national_rank, 'total': national_total_rank},
                'school': {'rank': school_rank, 'total': school_total_rank},
                'grade': {'rank': grade_rank, 'total': grade_total},
            },
            'statistics': {
                'national_average': round(float(national_average), 1) if national_average is not None else 0,
                'school_average': round(float(school_average), 1) if school_average else 0,
                'school_highest': school_high,
                'school_lowest': school_low,
            },
            'question_details': question_details,
            'comment': result.comment or '',
        }
        subject_entries.append(subject_code)

    if not subject_entries:
        return None, '成績データが見つかりません'

    subject_entries.sort(key=lambda code: SUBJECT_ORDER.get(code, 99))

    total_score = sum(subjects_data[code]['total_score'] for code in subject_entries)
    total_max = sum(subjects_data[code]['max_score'] for code in subject_entries)
    combined_metrics = _calculate_combined_metrics(student, year, period, total_score)

    report_data = {
        'student_info': student_info,
        'test_info': test_info,
        'subjects': subjects_data,
        'subject_order': subject_entries,
        'combined': {
            'total_score': total_score,
            'max_score': total_max,
            'rankings': {
                'grade': {
                    'rank': combined_metrics['grade_rank'],
                    'total': combined_metrics['grade_total']
                },
                'school': {
                    'rank': combined_metrics['school_rank'],
                    'total': combined_metrics['school_total']
                },
                'national': {
                    'rank': combined_metrics['national_rank'],
                    'total': combined_metrics['national_total']
                },
            },
            'averages': {
                'grade': combined_metrics['grade_average'],
                'school': combined_metrics['school_average'],
                'national': combined_metrics['national_average'],
            },
            'deviations': {
                'grade': combined_metrics['grade_deviation'],
                'school': combined_metrics['school_deviation'],
                'national': combined_metrics['national_deviation'],
            }
        },
        'trend': _collect_trend_data(student, grade_level),
    }

    return report_data, None


def _format_rank(rank_info: dict) -> str:
    if not rank_info or not rank_info.get('rank'):
        return '-'
    rank = rank_info['rank']
    total = rank_info.get('total')
    if total:
        return f"{int(rank)}位/{int(total)}人"
    return f"{int(rank)}位"


def _format_score(value) -> str:
    if value is None:
        return '-'
    if isinstance(value, (int, float, Decimal)):
        if float(value).is_integer():
            return f"{int(value)}"
        return f"{float(value):.1f}"
    return str(value)


def _generate_trend_chart(points: list[dict], title: str) -> str | None:
    if len(points) < 2:
        return None
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        matplotlib.rcParams['font.family'] = [
            'Noto Sans CJK JP', 'Yu Gothic', 'Hiragino Sans', 'Meiryo',
            'IPAexGothic', 'TakaoGothic', 'MS Gothic', 'DejaVu Sans'
        ]
        labels = [p['label'] for p in points]
        scores = [p['score'] for p in points]
        averages = [p.get('average', 0) for p in points]
        fig, ax = plt.subplots(figsize=(3.3, 2.0), dpi=160)
        ax.plot(labels, scores, marker='o', linewidth=2, label='あなた')
        if any(averages):
            ax.plot(labels, averages, marker='o', linestyle='--', linewidth=1.6, label='平均')
        ax.set_title(title)
        ax.grid(True, alpha=0.3)
        ax.set_ylim(bottom=0)
        ax.legend(fontsize=8, loc='lower right')
        fig.tight_layout()
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        fig.savefig(temp_file.name, transparent=True)
        plt.close(fig)
        return temp_file.name
    except Exception:
        return None


def create_individual_report_pdf(report_data: dict) -> tuple[str | None, str | None]:
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import landscape, A4
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    except ImportError as exc:
        return None, f'PDF生成に必要なライブラリが不足しています: {exc}'

    _register_pdf_fonts()

    reports_dir = _ensure_reports_dir()
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'reports', 'logo.png')

    file_name = "individual_report_{sid}_{year}_{period}_{stamp}.pdf".format(
        sid=report_data['student_info']['id'],
        year=report_data['test_info']['year'],
        period=report_data['test_info']['period'],
        stamp=datetime.now().strftime('%Y%m%d_%H%M%S'),
    )
    file_path = os.path.join(reports_dir, file_name)

    canvas_obj = canvas.Canvas(file_path, pagesize=landscape(A4))
    width, height = landscape(A4)

    # ヘッダー背景
    header_height = 40 * mm
    canvas_obj.setFillColor(colors.HexColor('#0b1f2c'))
    canvas_obj.rect(0, height - header_height, width, header_height, stroke=0, fill=1)

    canvas_obj.setFillColor(colors.white)
    canvas_obj.setFont('HeiseiKakuGo-W5', 22)
    canvas_obj.drawString(20 * mm, height - 18 * mm, '全国学力向上テスト 個人成績表')

    if os.path.exists(logo_path):
        try:
            logo_width = 42 * mm
            logo_height = 28 * mm
            canvas_obj.drawImage(
                logo_path,
                width - logo_width - 15 * mm,
                height - logo_height - 10 * mm,
                width=logo_width,
                height=logo_height,
                mask='auto'
            )
        except Exception:
            pass

    test_info = report_data['test_info']
    header_text = f"{test_info['year']}年度 {test_info['iteration']} {test_info['period_display']}"
    canvas_obj.setFont('HeiseiKakuGo-W5', 13)
    canvas_obj.drawString(20 * mm, height - 28 * mm, header_text.strip())

    canvas_obj.setFont('HeiseiKakuGo-W5', 10)
    canvas_obj.drawRightString(width - 15 * mm, height - 12 * mm, f"発行日: {datetime.now():%Y.%m.%d}")

    student_info = report_data['student_info']
    info_lines = [
        f"生徒名: {student_info['name']}",
        f"生徒ID: {student_info['id']}",
        f"学年: {student_info['grade']}   塾: {student_info['school_name']}   教室: {student_info['classroom_name']}",
    ]
    base_y = height - header_height - 8 * mm
    canvas_obj.setFillColor(colors.black)
    canvas_obj.setFont('HeiseiKakuGo-W5', 11)
    for idx, line in enumerate(info_lines):
        canvas_obj.drawString(20 * mm, base_y - idx * 6 * mm, line)

    # サマリーテーブル
    subjects_order = report_data['subject_order']
    column_headers = ['教科'] + [report_data['subjects'][code]['name'] for code in subjects_order] + ['合計']

    combined = report_data['combined']
    combined_rankings = combined['rankings']
    combined_averages = combined['averages']

    summary_rows = [column_headers]

    def _subject_row(key: str, formatter):
        row = [key]
        for code in subjects_order:
            row.append(formatter(report_data['subjects'][code]))
        row.append(formatter(combined))
        return row

    summary_rows.append(_subject_row('得点', lambda data: _format_score(data.get('total_score'))))
    summary_rows.append(_subject_row('偏差値', lambda data: _format_score(data.get('deviation'))))
    summary_rows.append(_subject_row('全国順位', lambda data: _format_rank(data.get('rankings', {}).get('national'))))
    summary_rows.append(_subject_row('塾内順位', lambda data: _format_rank(data.get('rankings', {}).get('school'))))
    summary_rows.append(_subject_row('平均点(全国)', lambda data: _format_score(data.get('statistics', {}).get('national_average') if 'statistics' in data else combined_averages.get('national'))))
    summary_rows.append(_subject_row('平均点(塾)', lambda data: _format_score(data.get('statistics', {}).get('school_average') if 'statistics' in data else combined_averages.get('school'))))
    summary_rows.append(_subject_row('最高点(塾)', lambda data: _format_score(data.get('statistics', {}).get('school_highest') if 'statistics' in data else None)))
    summary_rows.append(_subject_row('最低点(塾)', lambda data: _format_score(data.get('statistics', {}).get('school_lowest') if 'statistics' in data else None)))

    summary_table = Table(summary_rows, colWidths=[32 * mm] + [28 * mm] * (len(column_headers) - 1))
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0b1f2c')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'HeiseiKakuGo-W5'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f6f7f9')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#f6f7f9'), colors.HexColor('#ffffff')]),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
    ]))

    table_width, table_height = summary_table.wrap(0, 0)
    summary_x = 20 * mm
    summary_y = base_y - 45 * mm
    summary_table.drawOn(canvas_obj, summary_x, summary_y - table_height)

    # 科目別詳細
    detail_x = summary_x + table_width + 12 * mm
    detail_top = base_y - 10 * mm

    styles = getSampleStyleSheet()
    remark_style = ParagraphStyle(
        'Comment',
        parent=styles['Normal'],
        fontName='HeiseiKakuGo-W5',
        fontSize=9,
        leading=11,
    )

    current_top = detail_top
    for code in subjects_order:
        subject = report_data['subjects'][code]
        title_text = f"{subject['name']} 出題項目別の成果"
        canvas_obj.setFont('HeiseiKakuGo-W5', 12)
        canvas_obj.setFillColor(colors.HexColor('#0b1f2c'))
        canvas_obj.drawString(detail_x, current_top, title_text)

        question_rows = [['大問', '出題領域名', '得点/配点', '全国平均', '正答率']]
        for item in subject['question_details']:
            score_text = '-' if item['score'] is None else f"{item['score']}/{item['max_score']}"
            avg_text = f"{item['national_average']:.1f}" if item['national_average'] is not None else '-'
            rate_text = '-' if item['correct_rate'] is None else f"{item['correct_rate']:.0f}%"
            question_rows.append([
                f"{item['number']}",
                item['title'],
                score_text,
                avg_text,
                rate_text,
            ])

        question_table = Table(question_rows, colWidths=[18 * mm, 50 * mm, 28 * mm, 24 * mm, 24 * mm])
        question_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e1f0f8')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0b1f2c')),
            ('FONTNAME', (0, 0), (-1, -1), 'HeiseiKakuGo-W5'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#bcd4e6')),
        ]))

        qt_width, qt_height = question_table.wrap(0, 0)
        current_top -= 8 * mm
        question_table.drawOn(canvas_obj, detail_x, current_top - qt_height)
        current_top = current_top - qt_height - 8 * mm

        if subject['comment']:
            Paragraph(f"{subject['comment']}", remark_style).wrapOn(canvas_obj, qt_width, 20 * mm)
            Paragraph(f"{subject['comment']}", remark_style).drawOn(canvas_obj, detail_x, current_top - 18 * mm)
            current_top -= 22 * mm

    # 成長の推移チャート
    charts_top = summary_y - table_height - 20 * mm
    trend = report_data['trend']
    chart_x = 20 * mm

    overall_chart = _generate_trend_chart(trend.get('overall', []), '全教科の推移')
    if overall_chart:
        canvas_obj.drawImage(overall_chart, chart_x, charts_top - 55 * mm, width=65 * mm, height=40 * mm, mask='auto')
        try:
            os.remove(overall_chart)
        except OSError:
            pass

    chart_x += 70 * mm
    for code in subjects_order:
        chart_path = _generate_trend_chart(trend['subjects'].get(code, []), f"{report_data['subjects'][code]['name']}の推移")
        if not chart_path:
            continue
        canvas_obj.drawImage(chart_path, chart_x, charts_top - 55 * mm, width=65 * mm, height=40 * mm, mask='auto')
        chart_x += 70 * mm
        try:
            os.remove(chart_path)
        except OSError:
            pass

    canvas_obj.showPage()
    canvas_obj.save()

    try:
        os.chmod(file_path, 0o644)
    except OSError:
        pass

    return file_path, None


def generate_individual_report_template(student_id: str, year: int, period: str, format_type: str = 'pdf') -> dict:
    report_data, error = _collect_individual_report_data(student_id, year, period)
    if error:
        return {'success': False, 'error': error}

    if format_type == 'pdf':
        file_path, err = create_individual_report_pdf(report_data)
        if err:
            return {'success': False, 'error': err}
        download_url = _build_download_url(file_path)
        return {
            'success': True,
            'download_url': download_url,
            'format': 'pdf',
        }
    elif format_type == 'word':
        try:
            file_path = create_beautiful_word_report(report_data)
            download_url = _build_download_url(file_path)
            return {
                'success': True,
                'download_url': download_url,
                'format': 'word',
            }
        except Exception as exc:
            return {'success': False, 'error': str(exc)}
    else:
        return {'success': False, 'error': f'未対応の出力形式です: {format_type}'}


def generate_bulk_reports_template(student_ids: list[str], year: int, period: str, format_type: str = 'pdf') -> dict:
    if not student_ids:
        return {'success': False, 'error': '生徒IDが指定されていません'}

    generated_files = []
    errors = []

    for student_id in student_ids:
        report_data, data_error = _collect_individual_report_data(student_id, year, period)
        if data_error:
            errors.append(f"{student_id}: {data_error}")
            continue

        try:
            if format_type == 'pdf':
                file_path, err = create_individual_report_pdf(report_data)
                if err:
                    errors.append(f"{student_id}: {err}")
                    continue
            elif format_type == 'word':
                file_path = create_beautiful_word_report(report_data)
            else:
                errors.append(f"{student_id}: 未対応の出力形式です ({format_type})")
                continue

            generated_files.append(file_path)
        except Exception as exc:  # noqa: BLE001
            errors.append(f"{student_id}: {exc}")

    if not generated_files:
        return {'success': False, 'error': errors[0] if errors else '帳票生成に失敗しました'}

    reports_dir = _ensure_reports_dir()
    zip_name = "individual_reports_{year}_{period}_{stamp}.zip".format(
        year=year,
        period=period,
        stamp=datetime.now().strftime('%Y%m%d_%H%M%S')
    )
    zip_path = os.path.join(reports_dir, zip_name)

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for file_path in generated_files:
            arcname = os.path.basename(file_path)
            zip_file.write(file_path, arcname)

    try:
        os.chmod(zip_path, 0o644)
    except OSError:
        pass

    download_url = _build_download_url(zip_path)
    response = {
        'success': True,
        'download_url': download_url,
        'format': format_type,
    }

    if errors:
        response['warnings'] = errors

    return response
